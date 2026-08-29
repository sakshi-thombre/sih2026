"""Document text extraction.

Each loader takes raw bytes (never a filesystem path — the API layer
only ever hands us the uploaded file's content) and returns normalized
text grouped by page where the format supports it. File type is
determined by inspecting content, not trusted from the filename alone.

Never log document content here — only lengths/counts/types.
"""

import io
import re
from pathlib import Path

from docx import Document as DocxDocument
from pydantic import BaseModel
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.core.exceptions import InvalidDocumentError

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}

_TEXT_ENCODINGS = ("utf-8", "utf-16", "latin-1")


class ExtractedPage(BaseModel):
    page_number: int | None
    text: str


class ExtractedDocument(BaseModel):
    pages: list[ExtractedPage]
    file_type: str


def _clean_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _load_pdf(content: bytes) -> ExtractedDocument:
    if not content.startswith(b"%PDF"):
        raise InvalidDocumentError("File does not appear to be a valid PDF")
    try:
        reader = PdfReader(io.BytesIO(content))
        pages = [
            ExtractedPage(page_number=index, text=_clean_text(page.extract_text() or ""))
            for index, page in enumerate(reader.pages, start=1)
        ]
    except (PdfReadError, ValueError, KeyError) as exc:
        # pypdf raises several distinct exception types for corrupt/
        # malformed PDFs; all of them mean "we can't safely extract text".
        raise InvalidDocumentError("Could not read PDF file — it may be corrupted") from exc
    return ExtractedDocument(pages=pages, file_type="pdf")


def _load_docx(content: bytes) -> ExtractedDocument:
    if not content.startswith(b"PK"):
        raise InvalidDocumentError("File does not appear to be a valid DOCX")
    try:
        document = DocxDocument(io.BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    except Exception as exc:
        # python-docx doesn't expose a narrow exception type for corrupt
        # zip/xml content, so we convert any parse failure here.
        raise InvalidDocumentError("Could not read DOCX file — it may be corrupted") from exc
    # DOCX has no reliable notion of "page" until rendered, so we don't
    # invent one — page_number stays None per the citation contract.
    return ExtractedDocument(pages=[ExtractedPage(page_number=None, text=_clean_text(text))], file_type="docx")


def _load_txt(content: bytes) -> ExtractedDocument:
    for encoding in _TEXT_ENCODINGS:
        try:
            text = content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise InvalidDocumentError("Could not decode text file with a supported encoding")
    return ExtractedDocument(pages=[ExtractedPage(page_number=None, text=_clean_text(text))], file_type="txt")


def load_document(filename: str, content: bytes) -> ExtractedDocument:
    """Extract normalized text from an uploaded file's raw bytes.

    Raises `InvalidDocumentError` for empty, unsupported, corrupt, or
    text-less documents rather than letting extraction errors surface
    as unhandled exceptions.
    """
    if not content:
        raise InvalidDocumentError("Uploaded file is empty")

    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise InvalidDocumentError(f"Unsupported file type '{extension or 'unknown'}'")

    if extension == ".pdf":
        document = _load_pdf(content)
    elif extension == ".docx":
        document = _load_docx(content)
    else:
        document = _load_txt(content)

    if not any(page.text.strip() for page in document.pages):
        raise InvalidDocumentError("No extractable text found in document")

    return document
